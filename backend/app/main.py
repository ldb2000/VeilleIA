import os
import json
import re
import shutil
import subprocess
import time
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeoutError
from html import escape, unescape
from pathlib import Path
from fastapi import FastAPI, Depends, HTTPException, Response, Security
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPAuthorizationCredentials, HTTPBearer
from sqlalchemy.orm import Session
from sqlalchemy import inspect, text
from typing import List, Optional, Tuple
from datetime import datetime
from io import BytesIO

import google.generativeai as genai
import httpx
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from dotenv import load_dotenv
import markdown
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Image as RLImage, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from . import models, schemas, database
from .database import SessionLocal, engine

load_dotenv()

models.Base.metadata.create_all(bind=engine)


def ensure_reports_schema() -> None:
    inspector = inspect(engine)
    columns = {column["name"] for column in inspector.get_columns("reports")}
    with engine.begin() as connection:
        if "report_type" not in columns:
            connection.execute(
                text("ALTER TABLE reports ADD COLUMN report_type VARCHAR DEFAULT 'technical_watch'")
            )
            connection.execute(
                text("UPDATE reports SET report_type = 'technical_watch' WHERE report_type IS NULL")
            )


ensure_reports_schema()


def serialize_report_ids(report_ids: Optional[list[int]]) -> Optional[str]:
    if not report_ids:
        return None
    return json.dumps(report_ids)


def deserialize_report_ids(report_ids_raw: Optional[str]) -> Optional[list[int]]:
    if not report_ids_raw:
        return None
    try:
        value = json.loads(report_ids_raw)
    except json.JSONDecodeError:
        return None
    return value if isinstance(value, list) else None


def build_report_chat_context(db: Session, report: models.Report) -> str:
    summary = db.query(models.ReportSummary).filter(models.ReportSummary.report_id == report.id).first()
    if summary and summary.content.strip():
        body = f"Résumé exécutif enregistré :\n{summary.content.strip()}"
    else:
        compact_content = report.content.strip()
        if len(compact_content) > 3500:
            compact_content = compact_content[:3500].rstrip() + "\n\n[contenu tronqué]"
        body = f"Contenu du rapport :\n{compact_content}"

    return f"Rapport {report.id} | {report.date} | {report.created_at.isoformat()}\n{body}"

DEFAULT_GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.0-flash-lite")
DEFAULT_GEMINI_TIMEOUT_SECONDS = float(os.getenv("GEMINI_TIMEOUT_SECONDS", "90"))
REPORT_GEMINI_TIMEOUT_SECONDS = float(
    os.getenv("GEMINI_REPORT_TIMEOUT_SECONDS", str(max(DEFAULT_GEMINI_TIMEOUT_SECONDS, 300)))
)
GEMINI_EXECUTOR = ThreadPoolExecutor(max_workers=4)
CLI_MODEL_ALIASES = {
    "gemini-2.0-flash-lite": "auto-gemini-3",
    "gemini-1.5-flash": "auto-gemini-3",
    "gemini-2.0-flash": "auto-gemini-3",
}
MARKDOWN_IMAGE_RE = re.compile(r"!\[(?P<alt>.*?)\]\((?P<url>https?://[^\s)]+)\)")
MARKDOWN_LINK_RE = re.compile(r"\[(?P<label>[^\]]+)\]\((?P<url>https?://[^\s)]+)\)")
MARKDOWN_BOLD_RE = re.compile(r"\*\*(?P<text>.+?)\*\*")


def extract_gemini_cli_oauth_client_config() -> Tuple[Optional[str], Optional[str]]:
    gemini_path = shutil.which("gemini")
    if not gemini_path:
        return None, None

    try:
        resolved = Path(gemini_path).resolve()
        candidates = [
            resolved.parent / "node_modules" / "@google" / "gemini-cli-core" / "dist" / "src" / "code_assist" / "oauth2.js",
            resolved.parent.parent / "node_modules" / "@google" / "gemini-cli-core" / "dist" / "src" / "code_assist" / "oauth2.js",
            resolved.parent.parent / "node_modules" / "@google" / "gemini-cli" / "node_modules" / "@google" / "gemini-cli-core" / "dist" / "src" / "code_assist" / "oauth2.js",
        ]
        oauth2_path = next((path for path in candidates if path.exists()), None)
        if oauth2_path is None:
            matches = list(resolved.parent.parent.rglob("oauth2.js"))
            oauth2_path = next((path for path in matches if "code_assist" in str(path)), None)
        if oauth2_path is None:
            return None, None

        content = oauth2_path.read_text()
        client_id_match = re.search(r"(\d+-[a-z0-9]+\.apps\.googleusercontent\.com)", content)
        client_secret_match = re.search(r"(GOCSPX-[A-Za-z0-9_-]+)", content)
        client_id = client_id_match.group(1) if client_id_match else None
        client_secret = client_secret_match.group(1) if client_secret_match else None
        return client_id, client_secret
    except Exception:
        return None, None


def load_oauth_credentials(creds_path: Path) -> tuple[Credentials, dict]:
    with creds_path.open("r") as f:
        data = json.load(f)

    extracted_client_id, extracted_client_secret = extract_gemini_cli_oauth_client_config()
    client_id = (
        data.get("client_id")
        or os.getenv("GEMINI_OAUTH_CLIENT_ID")
        or extracted_client_id
    )
    client_secret = (
        data.get("client_secret")
        or os.getenv("GEMINI_OAUTH_CLIENT_SECRET")
        or extracted_client_secret
    )
    scopes = data.get("scopes") or data.get("scope")
    if isinstance(scopes, str):
        scopes = scopes.split()

    creds = Credentials(
        token=data.get("access_token") or data.get("token"),
        refresh_token=data.get("refresh_token"),
        token_uri=data.get("token_uri") or "https://oauth2.googleapis.com/token",
        client_id=client_id,
        client_secret=client_secret,
        scopes=scopes,
        id_token=data.get("id_token"),
    )

    expiry_date = data.get("expiry_date")
    if expiry_date:
        creds.expiry = datetime.utcfromtimestamp(expiry_date / 1000)

    data["client_id"] = client_id
    data["client_secret"] = client_secret
    return creds, data


def persist_oauth_credentials(creds_path: Path, source_data: dict, creds: Credentials) -> None:
    updated = dict(source_data)
    updated["access_token"] = creds.token
    if creds.refresh_token:
        updated["refresh_token"] = creds.refresh_token
    if creds.id_token:
        updated["id_token"] = creds.id_token
    if creds.client_id:
        updated["client_id"] = creds.client_id
    if creds.client_secret:
        updated["client_secret"] = creds.client_secret
    if creds.scopes:
        updated["scope"] = " ".join(creds.scopes)
    if creds.expiry:
        updated["expiry_date"] = int(creds.expiry.timestamp() * 1000)

    creds_path.parent.mkdir(parents=True, exist_ok=True)
    with creds_path.open("w") as f:
        json.dump(updated, f, indent=2)
        f.write("\n")


def load_gemini_auth():
    """
    Configures Gemini Developer API auth.
    Prefers ~/.gemini/oauth_creds.json when present, then falls back to GEMINI_API_KEY.
    """
    creds_path = Path.home() / ".gemini" / "oauth_creds.json"
    api_key = os.getenv("GEMINI_API_KEY")

    if creds_path.exists():
        if should_use_gemini_cli_backend():
            print(f"Using Gemini CLI backend with OAuth credentials from {creds_path}")
            return True
        try:
            creds, raw_creds = load_oauth_credentials(creds_path)

            if creds.refresh_token and creds.client_id and creds.client_secret:
                print("Refreshing Gemini OAuth access token")
                creds.refresh(Request())
                persist_oauth_credentials(creds_path, raw_creds, creds)
            elif not creds.token:
                raise RuntimeError(
                    "Gemini OAuth credentials are missing an access token. "
                    "Set GEMINI_OAUTH_CLIENT_ID and GEMINI_OAUTH_CLIENT_SECRET "
                    "to enable refresh, or re-authenticate with Gemini."
                )
            else:
                print(
                    "OAuth credentials loaded without backend refresh support. "
                    "Set GEMINI_OAUTH_CLIENT_ID and GEMINI_OAUTH_CLIENT_SECRET "
                    "to enable token refresh."
                )

            print(f"Configuring Gemini with OAuth credentials from {creds_path}")
            original_gemini_api_key = os.environ.pop("GEMINI_API_KEY", None)
            original_google_api_key = os.environ.pop("GOOGLE_API_KEY", None)
            try:
                genai.configure(credentials=creds)
            finally:
                if original_gemini_api_key is not None:
                    os.environ["GEMINI_API_KEY"] = original_gemini_api_key
                if original_google_api_key is not None:
                    os.environ["GOOGLE_API_KEY"] = original_google_api_key
            return True
        except Exception as e:
            print(f"Failed to load OAuth credentials: {e}")
            raise

    if api_key:
        print("Configuring Gemini with GEMINI_API_KEY")
        genai.configure(api_key=api_key)
        return True

    raise RuntimeError(
        "No Gemini authentication configured. Provide ~/.gemini/oauth_creds.json "
        "or set GEMINI_API_KEY."
    )


def get_auth_debug_info() -> dict:
    creds_path = Path.home() / ".gemini" / "oauth_creds.json"
    gemini_cli_path = shutil.which("gemini")
    info = {
        "has_oauth_file": creds_path.exists(),
        "oauth_path": str(creds_path),
        "has_gemini_api_key": bool(os.getenv("GEMINI_API_KEY")),
        "has_google_api_key": bool(os.getenv("GOOGLE_API_KEY")),
        "gemini_cli_path": gemini_cli_path,
        "use_cli_backend": bool(creds_path.exists() and gemini_cli_path),
        "model": DEFAULT_GEMINI_MODEL,
        "cli_model": resolve_cli_model_name(),
        "timeout_seconds": DEFAULT_GEMINI_TIMEOUT_SECONDS,
    }

    if creds_path.exists():
        creds, _ = load_oauth_credentials(creds_path)
        info.update(
            {
                "oauth_has_token": bool(creds.token),
                "oauth_has_refresh_token": bool(creds.refresh_token),
                "oauth_has_client_id": bool(creds.client_id),
                "oauth_has_client_secret": bool(creds.client_secret),
                "oauth_scopes": creds.scopes,
                "oauth_expiry": creds.expiry.isoformat() if creds.expiry else None,
                "oauth_valid": creds.valid,
                "oauth_expired": creds.expired,
            }
        )

    return info


def should_use_gemini_cli_backend() -> bool:
    return bool((Path.home() / ".gemini" / "oauth_creds.json").exists() and shutil.which("gemini"))


def resolve_cli_model_name() -> str:
    return os.getenv(
        "GEMINI_CLI_MODEL",
        CLI_MODEL_ALIASES.get(DEFAULT_GEMINI_MODEL, DEFAULT_GEMINI_MODEL),
    )


def generate_gemini_content_with_sdk(prompt: str, timeout_seconds: float = DEFAULT_GEMINI_TIMEOUT_SECONDS):
    model = genai.GenerativeModel(DEFAULT_GEMINI_MODEL)
    future = GEMINI_EXECUTOR.submit(
        model.generate_content,
        prompt,
        request_options={"timeout": timeout_seconds},
    )
    try:
        return future.result(timeout=timeout_seconds)
    except FuturesTimeoutError as e:
        future.cancel()
        raise TimeoutError(
            f"Gemini call exceeded application timeout of {timeout_seconds}s"
        ) from e


def generate_gemini_content_with_cli(
    prompt: str,
    timeout_seconds: float = DEFAULT_GEMINI_TIMEOUT_SECONDS,
) -> str:
    cli_model = resolve_cli_model_name()
    cmd = [
        "gemini",
        "--model",
        cli_model,
        "--prompt",
        prompt,
        "--output-format",
        "json",
    ]
    try:
        completed = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=timeout_seconds,
            check=False,
        )
    except subprocess.TimeoutExpired as e:
        raise TimeoutError(
            f"gemini CLI exceeded timeout of {timeout_seconds}s"
        ) from e
    stdout = (completed.stdout or "").strip()
    stderr = (completed.stderr or "").strip()

    if completed.returncode != 0:
        raise RuntimeError(
            f"gemini CLI failed with exit code {completed.returncode}: {stderr or stdout or 'no output'}"
        )
    if not stdout:
        raise RuntimeError(f"gemini CLI returned no output. stderr={stderr or '<empty>'}")

    try:
        payload = json.loads(stdout)
    except json.JSONDecodeError as e:
        raise RuntimeError(f"gemini CLI returned non-JSON output: {stdout[:500]}") from e

    if payload.get("error"):
        raise RuntimeError(f"gemini CLI error: {payload['error']}")

    response = payload.get("response")
    if not isinstance(response, str) or not response.strip():
        raise RuntimeError(f"gemini CLI JSON missing 'response': {stdout[:500]}")

    return response


def generate_gemini_content(prompt: str, timeout_seconds: float = DEFAULT_GEMINI_TIMEOUT_SECONDS):
    if should_use_gemini_cli_backend():
        return generate_gemini_content_with_cli(prompt, timeout_seconds=timeout_seconds)
    return generate_gemini_content_with_sdk(prompt, timeout_seconds=timeout_seconds)

app = FastAPI(title="AI Watch API")

ALLOWED_ORIGINS = os.getenv("ALLOWED_ORIGINS", "http://localhost:5173").split(",")

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


API_TOKEN = os.getenv("API_TOKEN")
security = HTTPBearer(auto_error=False)


def verify_token(credentials: Optional[HTTPAuthorizationCredentials] = Security(security)):
    if not API_TOKEN:
        return
    if not credentials or credentials.credentials != API_TOKEN:
        raise HTTPException(status_code=401, detail="Invalid or missing token")


TECHNICAL_PROMPT = """Recherche les nouveautés IA publiées dans les dernières 24h dans les catégories suivantes :

## 1. Labs / model providers
OpenAI, Anthropic, Google, Meta, Mistral, Microsoft, AWS, xAI, Cohere, NVIDIA

## 2. AI coding / vibe coding / dev agents / open source tools
OpenClaw, Paperclip, Cline, Continue, Aider, Void, OpenHands, OpenCode, Roo Code, Bolt.new, Codeium, Windsurf, Cursor, Claude Code, Gemini CLI, aider-like tools, terminal coding agents, self-hosted coding assistants

## 3. Écosystème connexe à surveiller
- MCP (Model Context Protocol)
- tool use / function calling
- browse / search / computer use
- coding agents
- IDE plugins
- VS Code / JetBrains AI plugins
- local-first AI tools
- self-hosted AI assistants
- Ollama ecosystem
- OpenRouter ecosystem
- agent orchestration frameworks
- sandboxing / permissions / security for agents
- memory / long context / retrieval for agents

Objectif :
Je veux une veille technique utile pour un profil architecture / plateforme / infra / sécurité / développement.
Je veux détecter rapidement les sorties réellement importantes, y compris les outils open source émergents qui peuvent changer les usages de développement, de prototypage ou de "vibe coding".

Consignes strictes :
- Ignore le bruit marketing autant que possible
- Priorise les sources officielles, changelogs, documentation, releases GitHub, notes de version, billets engineering
- Utilise les médias généralistes seulement en source secondaire
- Sépare clairement les faits confirmés des suppositions ou annonces vagues
- Ne retiens que les informations ayant un impact technique réel
- Si une annonce est purement marketing sans impact concret, indique-le brièvement puis écarte-la
- Mets en avant les nouveautés réellement disponibles, pas seulement annoncées
- Quand un projet open source monte vite, signale-le même s’il est encore immature
- Indique explicitement quand quelque chose est : prototype / alpha / bêta / GA / production-ready / expérimental

Analyse uniquement :
- nouveaux modèles
- nouvelles capacités des modèles
- changements d’API
- changements SDK / CLI / plugins
- tool use / search / browse / computer use
- support MCP / connecteurs / plugins
- multimodalité utile
- mémoire / long context / retrieval
- agents / orchestration / planification / exécution
- self-hosting / local-first / BYOK
- compatibilité OpenAI / Anthropic / Gemini / OpenRouter / Ollama / Azure OpenAI / modèles locaux
- sécurité / sandbox / permissions / secrets / isolation
- pricing uniquement si l’impact technique ou d’architecture est structurant
- breaking changes / migrations / dépréciations
- disponibilité réelle
- niveau de maturité réel
- cas d’usage concrets pour équipes techniques

Format obligatoire :

## 1. Nouveautés majeures
- 10 éléments maximum
- pour chaque élément :
  - ce qui sort
  - pourquoi c’est important techniquement
  - pour qui c’est utile
  - niveau de maturité
  - action recommandée

## 2. Détail par catégorie

### A. Labs / modèles
Pour chaque acteur concerné :
- nouveautés confirmées
- ce que cela change techniquement
- impacts pour dev / plateforme / infra / sécurité
- disponibilité réelle
- niveau de maturité
- action recommandée

### B. Outils AI coding / vibe coding / open source
Pour chaque outil concerné :
- nouveauté confirmée
- type d’outil (IDE, plugin, CLI, agent, orchestration, framework, self-hosted app)
- ce que cela change concrètement
- compatibilité modèles / providers
- mode de déploiement (cloud, local, self-hosted, hybride)
- risques / limites / sécurité
- niveau de maturité
- action recommandée

### C. Écosystème agents / MCP / tooling
- nouveautés confirmées
- connecteurs ou protocoles importants
- implications architecture / sécurité / gouvernance
- niveau de maturité
- action recommandée

## 3. Tableau de synthèse
Colonnes :
- catégorie
- acteur / outil
- nouveauté
- type
- impact technique
- maturité
- disponibilité
- action

## 4. À surveiller
- signaux faibles
- projets émergents
- repos GitHub qui montent vite
- nouveautés encore immatures mais potentiellement structurantes
- sujets à recontrôler dans les prochains jours

## 5. Filtre décisionnel
Termine par 3 sections courtes :
- À tester tout de suite
- À surveiller
- À ignorer pour l’instant
"""


CODIR_NOTE_PROMPT = """# RÔLE
Tu es Tech Lead à la STTI d'APRIL (assurance, Lyon). Tu produis une note 
hebdomadaire pour le CODIR à partir d'une veille technique IA détaillée.

# AUDIENCE
Directeurs APRIL — non-techniques, pressés, lecture en diagonale. Ils 
décident sur la base : risque / opportunité / coût.

# TÂCHE
À partir du document de veille fourni, produis une note CODIR au format 
Word (.docx), strictement 1 page A4.

# RÈGLES DE TRADUCTION (technique → business)
- **Inverse la pyramide** : impact métier d'abord, techno ensuite.
- **Supprime** : scores benchmarks (SWE-bench…), tailles de modèles 
  (128B, 256k context), noms de frameworks (Temporal, LangChain…), 
  rubriques purement marketing.
- **Traduis les acronymes** : RCE → "exécution de code à distance", 
  MCP → "connecteur universel de l'IA", etc.
- Chaque info doit répondre à : *"qu'est-ce que ça change pour APRIL ? 
  quel bénéfice ? quel risque d'inaction ?"*
- Ajoute une **deadline explicite** sur toute alerte de sécurité.

# STRUCTURE (dans cet ordre, rien de plus)
1. **En-tête** : "NOTE CODIR — VEILLE IA" + date à droite + sous-titre 
   "STTI / Architecture & Gouvernance IA"
2. **Synthèse** : un seul paragraphe, 4 lignes max, mots-clés en gras, 
   alerte sécurité en rouge
3. **Décisions proposées** : tableau de 3 lignes max (colonnes : # / 
   Action / Bénéfice attendu / Validateur). Ligne sécurité en rouge.
4. **Points de vigilance** : 3 bullets, recadrés sur l'impact APRIL 
   (mentionner notre stack quand pertinent : Dagster, Snowflake, Azure, 
   EasyVista, M365…)
5. **Pied de page** : rédacteur + renvoi au rapport complet

# CONTRAINTES VISUELLES
- A4, marges resserrées (haut/bas 0.75", gauche/droite 1")
- Police Calibri, bleu corporate #1F4E79, rouge alerte #C0504D
- Tableau : en-tête bleu plein, alternance gris clair / blanc

# CONTRAINTE DURE
1 page A4 stricte. Si dépassement, sacrifier d'abord les "à surveiller", 
puis raccourcir la synthèse. Jamais réduire le tableau de décisions.

# FORMAT DE SORTIE
Retourne uniquement le contenu structuré en Markdown simple. N'ajoute pas de bloc de code.
Le document Word sera généré ensuite par l'application.

# DOCUMENT DE VEILLE EN ENTRÉE
{watch_document}
"""


def decode_html_entities(text: str) -> str:
    decoded = text
    for _ in range(3):
        next_decoded = unescape(decoded)
        if next_decoded == decoded:
            break
        decoded = next_decoded
    return decoded


def markdown_to_paragraph_markup(text: str) -> str:
    escaped = escape(decode_html_entities(text), quote=False)
    escaped = MARKDOWN_LINK_RE.sub(
        lambda match: (
            f'<link href="{escape(decode_html_entities(match.group("url")), quote=True)}" color="#1d4ed8">'
            f'{escape(decode_html_entities(match.group("label")), quote=False)}</link>'
        ),
        escaped,
    )
    escaped = MARKDOWN_BOLD_RE.sub(
        lambda match: f"<b>{escape(decode_html_entities(match.group('text')), quote=False)}</b>",
        escaped,
    )
    return escaped.replace("\n", "<br/>")


def pdf_text(text: str) -> str:
    return escape(decode_html_entities(text), quote=False)


def build_pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="WatchTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=22,
            leading=28,
            textColor=colors.HexColor("#0f172a"),
            alignment=TA_LEFT,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="WatchMeta",
            parent=styles["Normal"],
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#475569"),
            alignment=TA_LEFT,
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="WatchBody",
            parent=styles["Normal"],
            fontSize=10.5,
            leading=15,
            textColor=colors.HexColor("#1e293b"),
            alignment=TA_LEFT,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="WatchBullet",
            parent=styles["Normal"],
            fontSize=10.5,
            leading=15,
            leftIndent=12,
            firstLineIndent=-8,
            bulletIndent=0,
            textColor=colors.HexColor("#1e293b"),
            alignment=TA_LEFT,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="WatchH2",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=15,
            leading=19,
            textColor=colors.HexColor("#0f172a"),
            alignment=TA_LEFT,
            spaceBefore=10,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="WatchCaption",
            parent=styles["Normal"],
            fontSize=8.5,
            leading=11,
            textColor=colors.HexColor("#64748b"),
            alignment=TA_LEFT,
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="WatchH3",
            parent=styles["Heading3"],
            fontName="Helvetica-Bold",
            fontSize=12,
            leading=15,
            textColor=colors.HexColor("#334155"),
            alignment=TA_LEFT,
            spaceBefore=8,
            spaceAfter=5,
        )
    )
    styles.add(
        ParagraphStyle(
            name="WatchQuote",
            parent=styles["Normal"],
            fontSize=9.5,
            leading=13,
            leftIndent=10,
            rightIndent=6,
            textColor=colors.HexColor("#475569"),
            alignment=TA_LEFT,
            spaceBefore=2,
            spaceAfter=6,
            borderColor=colors.HexColor("#cbd5e1"),
            borderWidth=0.75,
            borderPadding=6,
        )
    )
    return styles


def build_pdf_header(report_date: str, styles):
    header_table = Table(
        [
            [Paragraph("AI Technical Watch", styles["WatchTitle"])],
            [Paragraph(report_date, styles["WatchMeta"])],
        ],
        colWidths=[170 * mm],
    )
    header_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f8fafc")),
                ("BOX", (0, 0), (-1, -1), 0.75, colors.HexColor("#cbd5e1")),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 10),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    return header_table


def fetch_pdf_image(url: str):
    try:
        response = httpx.get(url, timeout=8.0, follow_redirects=True)
        if response.status_code >= 400:
            return None
        content_type = response.headers.get("content-type", "")
        if "image" not in content_type:
            return None
        image = RLImage(BytesIO(response.content))
        image._restrictSize(170 * mm, 90 * mm)
        return image
    except Exception:
        return None


def build_markdown_table(lines: list[str], styles):
    raw_rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        raw_rows.append(cells)

    if len(raw_rows) < 2:
        return []

    header = raw_rows[0]
    data_rows = raw_rows[2:] if len(raw_rows) > 2 else []
    table_data = [[Paragraph(markdown_to_paragraph_markup(cell), styles["WatchBody"]) for cell in header]]
    for row in data_rows:
        normalized = row + [""] * (len(header) - len(row))
        table_data.append(
            [Paragraph(markdown_to_paragraph_markup(cell), styles["WatchBody"]) for cell in normalized[: len(header)]]
        )

    table = Table(table_data, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e2e8f0")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0f172a")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    return [table, Spacer(1, 10)]


def append_markdown_pdf_elements(elements, content: str, styles) -> None:
    lines = content.splitlines()
    index = 0

    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()

        if not stripped:
            index += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            elements.extend(build_markdown_table(table_lines, styles))
            continue

        image_match = MARKDOWN_IMAGE_RE.fullmatch(stripped)
        if image_match:
            image = fetch_pdf_image(image_match.group("url"))
            if image is not None:
                elements.append(image)
                caption = image_match.group("alt").strip()
                if caption:
                    elements.append(Paragraph(pdf_text(caption), styles["WatchCaption"]))
            index += 1
            continue

        if stripped.startswith("## "):
            elements.append(Paragraph(pdf_text(stripped[3:]), styles["WatchH2"]))
            index += 1
            continue

        if stripped.startswith("- "):
            elements.append(
                Paragraph(markdown_to_paragraph_markup(stripped[2:]), styles["WatchBullet"], bulletText="•")
            )
            index += 1
            continue

        elements.append(Paragraph(markdown_to_paragraph_markup(stripped), styles["WatchBody"]))
        index += 1


def append_note_pdf_elements(elements, note: models.Note, styles) -> None:
    created_at = note.created_at.strftime("%Y-%m-%d %H:%M") if note.created_at else ""
    kind_labels = {
        "detail": "Détail",
        "definition": "Définition",
        "note": "Note",
    }
    title_parts = [f"{kind_labels.get(note.kind, note.kind or 'Note')} #{note.id}"]
    if created_at:
        title_parts.append(created_at)

    elements.append(Paragraph(pdf_text(" - ".join(title_parts)), styles["WatchH3"]))
    if note.source_text and note.source_text.strip():
        elements.append(Paragraph("<b>Extrait :</b>", styles["WatchCaption"]))
        elements.append(Paragraph(markdown_to_paragraph_markup(note.source_text.strip()), styles["WatchQuote"]))
    if note.content and note.content.strip():
        append_markdown_pdf_elements(elements, note.content.strip(), styles)


def build_pdf_story(
    content: str,
    report_date: str,
    summary: Optional[models.ReportSummary] = None,
    notes: Optional[List[models.Note]] = None,
):
    styles = build_pdf_styles()
    elements = [build_pdf_header(report_date, styles), Spacer(1, 12)]

    if summary and summary.content and summary.content.strip():
        elements.append(Paragraph("Résumé exécutif", styles["WatchH2"]))
        append_markdown_pdf_elements(elements, summary.content.strip(), styles)
        elements.append(Spacer(1, 8))

    elements.append(Paragraph("Rapport complet", styles["WatchH2"]))
    append_markdown_pdf_elements(elements, content, styles)

    report_notes = notes or []
    if report_notes:
        elements.append(Spacer(1, 10))
        elements.append(Paragraph("Notes, détails et définitions", styles["WatchH2"]))
        for note in report_notes:
            append_note_pdf_elements(elements, note, styles)

    return elements


def clean_markdown_text(text: str) -> str:
    cleaned = decode_html_entities(text).strip()
    cleaned = re.sub(r"^#{1,6}\s+", "", cleaned)
    cleaned = cleaned.replace("**", "")
    cleaned = cleaned.replace("__", "")
    cleaned = cleaned.replace("—", "-")
    return cleaned.strip()


def is_security_text(text: str) -> bool:
    lowered = text.lower()
    markers = [
        "sécurité",
        "securite",
        "alerte",
        "vulnérabilité",
        "vulnerabilite",
        "exécution de code",
        "execution de code",
        "critique",
        "rce",
    ]
    return any(marker in lowered for marker in markers)


def add_markdown_runs(paragraph, text: str, color=None, bold_default: bool = False) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", decode_html_entities(text))
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        run_text = part[2:-2] if is_bold else part
        run = paragraph.add_run(run_text)
        run.bold = bold_default or is_bold
        if color is not None:
            run.font.color.rgb = color


def set_docx_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def set_docx_cell_text(cell, text: str, color=None, bold: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.text = ""
    run = paragraph.add_run(clean_markdown_text(text))
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def parse_markdown_table(table_lines: List[str]) -> List[List[str]]:
    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if cells and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells):
            continue
        rows.append(cells)
    return rows


def append_docx_markdown(document, content: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    blue = RGBColor(0x1F, 0x4E, 0x79)
    red = RGBColor(0xC0, 0x50, 0x4D)
    lines = content.splitlines()
    index = 0

    while index < len(lines):
        raw_line = lines[index].rstrip()
        stripped = raw_line.strip()
        if not stripped:
            index += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            rows = parse_markdown_table(table_lines)
            if rows:
                table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
                table.style = "Table Grid"
                for row_index, row in enumerate(rows):
                    for col_index, value in enumerate(row):
                        cell = table.cell(row_index, col_index)
                        if row_index == 0:
                            set_docx_cell_shading(cell, "1F4E79")
                            set_docx_cell_text(cell, value, color=RGBColor(0xFF, 0xFF, 0xFF), bold=True)
                        else:
                            set_docx_cell_shading(cell, "F2F2F2" if row_index % 2 == 1 else "FFFFFF")
                            set_docx_cell_text(cell, value, color=red if is_security_text(" ".join(row)) else None)
                continue

        if stripped.startswith("#"):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(2)
            add_markdown_runs(paragraph, clean_markdown_text(stripped), color=blue, bold_default=True)
            index += 1
            continue

        if stripped.startswith(("- ", "* ")):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(1)
            add_markdown_runs(paragraph, stripped[2:], color=red if is_security_text(stripped) else None)
            index += 1
            continue

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(2)
        add_markdown_runs(paragraph, stripped, color=red if is_security_text(stripped) else None)
        index += 1


def build_codir_docx(report: models.Report) -> bytes:
    try:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as e:
        raise RuntimeError("python-docx is required to export CODIR notes. Run pip install -r requirements.txt") from e

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(8.8)

    blue = RGBColor(0x1F, 0x4E, 0x79)
    red = RGBColor(0xC0, 0x50, 0x4D)

    header_table = document.add_table(rows=2, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = True
    title_cell = header_table.cell(0, 0)
    title_run = title_cell.paragraphs[0].add_run("NOTE CODIR - VEILLE IA")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = blue

    date_cell = header_table.cell(0, 1)
    date_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_cell.paragraphs[0].add_run(report.date)
    date_run.bold = True
    date_run.font.color.rgb = blue

    subtitle_cell = header_table.cell(1, 0)
    subtitle_cell.merge(header_table.cell(1, 1))
    subtitle_run = subtitle_cell.paragraphs[0].add_run("STTI / Architecture & Gouvernance IA")
    subtitle_run.font.color.rgb = blue
    subtitle_run.font.size = Pt(9)

    body_content = "\n".join(
        line
        for line in report.content.splitlines()
        if "NOTE CODIR" not in line.upper()
        and "STTI / ARCHITECTURE" not in line.upper()
    )
    append_docx_markdown(document, body_content)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Rédacteur : STTI / Tech Lead - Renvoi au rapport complet de veille IA")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = red if is_security_text(report.content) else blue

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()

@app.get("/reports", response_model=List[schemas.Report])
def read_reports(skip: int = 0, limit: int = 100, db: Session = Depends(get_db), _=Depends(verify_token)):
    reports = db.query(models.Report).order_by(models.Report.created_at.desc()).offset(skip).limit(limit).all()
    return reports

@app.get("/reports/{report_id}", response_model=schemas.Report)
def read_report(report_id: int, db: Session = Depends(get_db), _=Depends(verify_token)):
    db_report = db.query(models.Report).filter(models.Report.id == report_id).first()
    if db_report is None:
        raise HTTPException(status_code=404, detail="Report not found")
    return db_report


@app.delete("/reports/{report_id}", status_code=204)
def delete_report(report_id: int, db: Session = Depends(get_db), _=Depends(verify_token)):
    db_report = db.query(models.Report).filter(models.Report.id == report_id).first()
    if db_report is None:
        raise HTTPException(status_code=404, detail="Report not found")
    db.query(models.ReportSummary).filter(models.ReportSummary.report_id == report_id).delete()
    db.query(models.Note).filter(models.Note.report_id == report_id).delete()
    db.delete(db_report)
    db.commit()
    return Response(status_code=204)


@app.get("/reports/{report_id}/summary", response_model=schemas.ReportSummaryResponse)
def get_report_summary(report_id: int, db: Session = Depends(get_db), _=Depends(verify_token)):
    db_summary = db.query(models.ReportSummary).filter(models.ReportSummary.report_id == report_id).first()
    if db_summary is None:
        raise HTTPException(status_code=404, detail="Summary not found")
    return {
        "report_id": db_summary.report_id,
        "summary": db_summary.content,
        "updated_at": db_summary.updated_at,
    }


@app.post("/reports/{report_id}/summary", response_model=schemas.ReportSummaryResponse)
def summarize_report(report_id: int, db: Session = Depends(get_db), _=Depends(verify_token)):
    db_report = db.query(models.Report).filter(models.Report.id == report_id).first()
    if db_report is None:
        raise HTTPException(status_code=404, detail="Report not found")

    summary_prompt = f"""Tu reçois un rapport de veille déjà rédigé.

Ta tâche : produire un executive summary court, sans refaire de recherche.

Contraintes :
- utilise uniquement le contenu fourni ci-dessous
- pas de nouvelles affirmations externes
- 5 points maximum
- ton très concret
- orienté architecture / plateforme / sécurité / delivery
- termine par une ligne `Décision suggérée : ...`

Rapport :

{db_report.content}
"""

    try:
        load_gemini_auth()
        response = generate_gemini_content(summary_prompt)
        content = response if isinstance(response, str) else response.text
        db_summary = db.query(models.ReportSummary).filter(models.ReportSummary.report_id == report_id).first()
        if db_summary is None:
            db_summary = models.ReportSummary(report_id=report_id, content=content)
            db.add(db_summary)
        else:
            db_summary.content = content
            db_summary.updated_at = datetime.utcnow()
        db.commit()
        db.refresh(db_summary)
        return {
            "report_id": db_summary.report_id,
            "summary": db_summary.content,
            "updated_at": db_summary.updated_at,
        }
    except Exception as e:
        print(f"Error generating summary: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/reports/{report_id}/detail")
def expand_report_selection(
    report_id: int,
    selected_text: str,
    db: Session = Depends(get_db),
    _=Depends(verify_token),
):
    db_report = db.query(models.Report).filter(models.Report.id == report_id).first()
    if db_report is None:
        raise HTTPException(status_code=404, detail="Report not found")

    detail_prompt = f"""Tu reçois un rapport de veille déjà rédigé et un extrait sélectionné par l'utilisateur.

Ta tâche : expliquer plus en détail cet extrait, sans refaire de recherche.

Contraintes :
- utilise uniquement le rapport ci-dessous
- ne crée pas de nouveaux faits externes
- 6 lignes maximum
- ton concret, technique, utile
- précise si le rapport ne contient pas assez d'information

Extrait sélectionné :
{selected_text}

Rapport complet :
{db_report.content}
"""

    try:
        load_gemini_auth()
        response = generate_gemini_content(detail_prompt)
        content = response if isinstance(response, str) else response.text
        return {"detail": content}
    except Exception as e:
        print(f"Error generating detail: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/reports/{report_id}/definition", response_model=schemas.ReportDefinitionResponse)
def define_report_selection(
    report_id: int,
    selected_text: str,
    db: Session = Depends(get_db),
    _=Depends(verify_token),
):
    db_report = db.query(models.Report).filter(models.Report.id == report_id).first()
    if db_report is None:
        raise HTTPException(status_code=404, detail="Report not found")

    definition_prompt = f"""Tu reçois un rapport de veille IA déjà rédigé et un terme ou extrait sélectionné par l'utilisateur.

Ta tâche : produire une définition courte, utile et exploitable qui pourra être ajoutée en fin de document.

Contraintes :
- explique le concept en français clair
- commence par une phrase de définition
- ajoute 2 à 3 points maximum : pourquoi c'est important, risque ou usage concret
- traduis les acronymes si nécessaire
- si le rapport ne donne pas assez de contexte, donne une définition générale et indique qu'elle est hors contexte rapport
- ne fais pas de recherche web en direct

Terme ou extrait sélectionné :
{selected_text}

Rapport complet :
{db_report.content}
"""

    try:
        load_gemini_auth()
        response = generate_gemini_content(definition_prompt)
        content = response if isinstance(response, str) else response.text
        return {"definition": content}
    except Exception as e:
        print(f"Error generating definition: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/notes", response_model=List[schemas.Note])
def read_notes(db: Session = Depends(get_db), _=Depends(verify_token)):
    return db.query(models.Note).order_by(models.Note.created_at.desc()).all()


@app.post("/reports/{report_id}/notes", response_model=schemas.Note)
def create_note(report_id: int, payload: schemas.NoteCreate, db: Session = Depends(get_db), _=Depends(verify_token)):
    db_report = db.query(models.Report).filter(models.Report.id == report_id).first()
    if db_report is None:
        raise HTTPException(status_code=404, detail="Report not found")

    db_note = models.Note(
        report_id=report_id,
        kind=payload.kind,
        source_text=payload.source_text,
        content=payload.content,
    )
    db.add(db_note)
    db.commit()
    db.refresh(db_note)
    return db_note


@app.delete("/notes/{note_id}", status_code=204)
def delete_note(note_id: int, db: Session = Depends(get_db), _=Depends(verify_token)):
    db_note = db.query(models.Note).filter(models.Note.id == note_id).first()
    if db_note is None:
        raise HTTPException(status_code=404, detail="Note not found")
    db.delete(db_note)
    db.commit()
    return Response(status_code=204)


@app.get("/chat/messages", response_model=List[schemas.ChatMessage])
def get_chat_messages(db: Session = Depends(get_db), _=Depends(verify_token)):
    messages = db.query(models.ChatMessage).order_by(models.ChatMessage.created_at.asc()).all()
    return [
        {
            "id": message.id,
            "role": message.role,
            "content": message.content,
            "report_ids": deserialize_report_ids(message.report_ids),
            "created_at": message.created_at,
        }
        for message in messages
    ]


@app.post("/chat", response_model=schemas.ChatResponse)
def chat_with_reports(payload: schemas.ChatRequest, db: Session = Depends(get_db), _=Depends(verify_token)):
    question = payload.question.strip()
    if not question:
        raise HTTPException(status_code=400, detail="Question cannot be empty")

    query = db.query(models.Report).order_by(models.Report.created_at.desc())
    if payload.report_ids:
      reports = query.filter(models.Report.id.in_(payload.report_ids)).limit(3).all()
    else:
      reports = query.limit(3).all()

    if not reports:
        raise HTTPException(status_code=404, detail="No reports available")

    user_message = models.ChatMessage(
        role="user",
        content=question,
        report_ids=serialize_report_ids(payload.report_ids),
    )
    db.add(user_message)
    db.commit()

    context_blocks = [build_report_chat_context(db, report) for report in reports]
    context_text = "\n\n---\n\n".join(context_blocks)

    chat_prompt = f"""Tu es un assistant technique qui répond à des questions à partir de rapports de veille IA déjà générés.

Règles :
- priorise les informations présentes dans les rapports fournis
- si la question demande un complément lié à la technologie, tu peux donner du contexte général utile, mais indique clairement ce qui vient des rapports et ce qui est du contexte général
- ne prétends pas avoir fait une recherche web en direct
- si l'information n'est pas présente, dis-le clairement
- réponse dense, concrète, orientée architecture / plateforme / sécurité / développement
- utilise des listes courtes si utile
- si plusieurs rapports se contredisent ou si le contexte est insuffisant, signale-le

Question :
{question}

Rapports de contexte :

{context_text}
"""

    try:
        load_gemini_auth()
        response = generate_gemini_content(chat_prompt)
        content = response if isinstance(response, str) else response.text
        db.add(
            models.ChatMessage(
                role="assistant",
                content=content,
                report_ids=serialize_report_ids(payload.report_ids),
            )
        )
        db.commit()
        return {"answer": content}
    except Exception as e:
        print(f"Error in chat endpoint: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/reports/trigger", response_model=schemas.Report)
async def trigger_report(db: Session = Depends(get_db), _=Depends(verify_token)):
    try:
        load_gemini_auth()
        backend_kind = "cli" if should_use_gemini_cli_backend() else "sdk"
        active_model = resolve_cli_model_name() if backend_kind == "cli" else DEFAULT_GEMINI_MODEL
        print(
            f"Sending request to Gemini backend={backend_kind} model={active_model} "
            f"timeout={REPORT_GEMINI_TIMEOUT_SECONDS}s"
        )
        start = time.perf_counter()
        response = generate_gemini_content(
            TECHNICAL_PROMPT,
            timeout_seconds=REPORT_GEMINI_TIMEOUT_SECONDS,
        )
        elapsed = time.perf_counter() - start
        print(f"Gemini response received in {elapsed:.2f}s")
        content = response if isinstance(response, str) else response.text
        
        db_report = models.Report(
            date=datetime.now().strftime("%Y-%m-%d"),
            report_type="technical_watch",
            content=content
        )
        db.add(db_report)
        db.commit()
        db.refresh(db_report)
        return db_report
    except Exception as e:
        print(f"Error triggering report: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/reports/{report_id}/codir", response_model=schemas.Report)
def generate_codir_note(report_id: int, db: Session = Depends(get_db), _=Depends(verify_token)):
    source_report = db.query(models.Report).filter(models.Report.id == report_id).first()
    if source_report is None:
        raise HTTPException(status_code=404, detail="Report not found")

    source_summary = db.query(models.ReportSummary).filter(models.ReportSummary.report_id == report_id).first()
    source_notes = (
        db.query(models.Note)
        .filter(models.Note.report_id == report_id)
        .order_by(models.Note.created_at.asc())
        .all()
    )

    watch_parts = [f"# Rapport complet - {source_report.date}", source_report.content]
    if source_summary and source_summary.content.strip():
        watch_parts.append(f"# Résumé exécutif enregistré\n{source_summary.content.strip()}")
    if source_notes:
        notes_text = "\n\n".join(
            f"- {note.kind}: {note.content.strip()}" for note in source_notes if note.content and note.content.strip()
        )
        if notes_text:
            watch_parts.append(f"# Notes, détails et définitions enregistrés\n{notes_text}")

    prompt = CODIR_NOTE_PROMPT.format(watch_document="\n\n".join(watch_parts))

    try:
        load_gemini_auth()
        backend_kind = "cli" if should_use_gemini_cli_backend() else "sdk"
        active_model = resolve_cli_model_name() if backend_kind == "cli" else DEFAULT_GEMINI_MODEL
        print(
            f"Sending CODIR note request to Gemini backend={backend_kind} model={active_model} "
            f"timeout={REPORT_GEMINI_TIMEOUT_SECONDS}s"
        )
        response = generate_gemini_content(prompt, timeout_seconds=REPORT_GEMINI_TIMEOUT_SECONDS)
        content = response if isinstance(response, str) else response.text

        db_report = models.Report(
            date=datetime.now().strftime("%Y-%m-%d"),
            report_type="codir_note",
            content=content,
        )
        db.add(db_report)
        db.commit()
        db.refresh(db_report)
        return db_report
    except Exception as e:
        print(f"Error generating CODIR note: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/reports/{report_id}/pdf")
def get_report_pdf(report_id: int, db: Session = Depends(get_db), _=Depends(verify_token)):
    db_report = db.query(models.Report).filter(models.Report.id == report_id).first()
    if db_report is None:
        raise HTTPException(status_code=404, detail="Report not found")
    db_summary = db.query(models.ReportSummary).filter(models.ReportSummary.report_id == report_id).first()
    db_notes = (
        db.query(models.Note)
        .filter(models.Note.report_id == report_id)
        .order_by(models.Note.created_at.asc())
        .all()
    )
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=f"AI Technical Watch - {db_report.date}",
    )
    elements = build_pdf_story(db_report.content, db_report.date, summary=db_summary, notes=db_notes)
    doc.build(elements)
    
    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f"attachment; filename=ai_watch_{db_report.date}.pdf"
        }
    )


@app.get("/reports/{report_id}/docx")
def get_report_docx(report_id: int, db: Session = Depends(get_db), _=Depends(verify_token)):
    db_report = db.query(models.Report).filter(models.Report.id == report_id).first()
    if db_report is None:
        raise HTTPException(status_code=404, detail="Report not found")
    if db_report.report_type != "codir_note":
        raise HTTPException(status_code=400, detail="DOCX export is only available for CODIR notes")

    docx_bytes = build_codir_docx(db_report)
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename=note_codir_veille_ia_{db_report.date}.docx"
        },
    )
